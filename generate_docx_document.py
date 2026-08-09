#!/usr/bin/env python3
"""
Generates CHARTORA_IN_MASTER_BUSINESS_MODEL_V1.docx (Official Word Document)
"""

import os
import sys

def create_word_docx():
    md_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'CHARTORA_IN_MASTER_BUSINESS_MODEL_V1.md')
    docx_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'CHARTORA_IN_MASTER_BUSINESS_MODEL_V1.docx')

    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        doc = docx.Document()
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue

            if line_str.startswith('# '):
                p = doc.add_heading(level=1)
                run = p.add_run(line_str.replace('# ', ''))
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(16, 185, 129) # Vibrant Emerald
            elif line_str.startswith('## '):
                p = doc.add_heading(level=2)
                run = p.add_run(line_str.replace('## ', ''))
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(4, 120, 87) # Rich Emerald
            elif line_str.startswith('### '):
                p = doc.add_heading(level=3)
                run = p.add_run(line_str.replace('### ', ''))
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(6, 78, 59)
            elif line_str.startswith('- ') or line_str.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                text = line_str[2:].replace('**', '')
                p.add_run(text)
            else:
                p = doc.add_paragraph()
                text = line_str.replace('**', '')
                p.add_run(text)

        doc.save(docx_file)
        print(f"✅ Successfully created Word Document: {docx_file}")

    except ImportError:
        print("python-docx not installed, creating fallback plain text Word document...")
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        with open(docx_file, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"✅ Fallback Word document saved to {docx_file}")

if __name__ == '__main__':
    create_word_docx()
